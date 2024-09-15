import logging

import markdown2
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.styles.style import BaseStyle, ParagraphStyle
from docx.text.parfmt import ParagraphFormat
from tqdm import tqdm
from docx.shared import RGBColor
import docx.oxml.ns


logger = logging.getLogger()
logger.setLevel(logging.DEBUG)


def update_style(
    doc, style_name: str, alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY, uppercase=False
):
    if style_name.startswith("Header"):
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[f"Header {style_name[-1]}"]

    style: ParagraphStyle = doc.styles[style_name]
    font = style.font
    font.color.rgb = RGBColor(0, 0, 0)
    font.name = "Times New Roman"
    font.size = Pt(14)
    font.bold = False
    font.italic = False

    # Заглавные буквы, если нужно
    if uppercase:
        font.all_caps = True

    paragraph_format: ParagraphFormat = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.5

    # Настройка выравнивания
    if alignment != WD_ALIGN_PARAGRAPH.CENTER:
        paragraph_format.first_line_indent = Cm(1.25)
    paragraph_format.alignment = alignment


def init_docx(docx_file):
    # Создаем новый документ
    doc = Document()

    # Создаем пользовательские стили для заголовков
    update_style(doc, "Normal")
    update_style(doc, "Header 1", WD_ALIGN_PARAGRAPH.CENTER, uppercase=True)
    update_style(doc, "Header 2", WD_ALIGN_PARAGRAPH.LEFT)
    update_style(doc, "Header 3", WD_ALIGN_PARAGRAPH.LEFT)
    update_style(doc, "Header 4", WD_ALIGN_PARAGRAPH.LEFT)

    return doc


def md_to_docx(md_file, docx_file):
    with open(md_file, "r", encoding="utf-8") as file:
        md_content = file.read()
    html_content = markdown2.markdown(md_content)
    doc = init_docx("styled_output.docx")
    logging.debug("Document initialized with styles")

    # Разбиваем HTML на блоки для обработки
    in_list = False
    for line in tqdm(html_content.splitlines(), ncols=80):
        if line.startswith("<h1>"):
            doc.add_paragraph(line[4:-5], style="Header 1")
            logger.debug(f"Added Header 1: {line[4:-5]}")
        elif line.startswith("<h2>"):
            doc.add_paragraph(line[4:-5], style="Header 2")
            logger.debug(f"Added Header 2: {line[4:-5]}")
        elif line.startswith("<h3>"):
            doc.add_paragraph(line[4:-5], style="Header 3")
            logger.debug(f"Added Header 3: {line[4:-5]}")
        elif line.startswith("<h4>"):
            doc.add_paragraph(line[4:-5], style="Header 4")
            logger.debug(f"Added Header 4: {line[4:-5]}")
        elif line.startswith("<ul>") or in_list:
            # Если начинается маркированный список
            in_list = True
            while in_list:
                line = line.strip()
                if line.startswith("<li>") or line.startswith("<ul>"):
                    doc.add_paragraph(line[4:-5], style="ListBullet")
                if "</ul>" in line or "</li>" in line:
                    in_list = False
                else:
                    break
        elif line.startswith("<p>"):
            # Добавляем новый абзац
            doc.add_paragraph(line[3:-4], style="Normal")

    # Сохраняем документ .docx
    doc.save(docx_file)


if __name__ == "__main__":
    # Пример использования
    md_to_docx("data/input.md", "data/output.docx")
